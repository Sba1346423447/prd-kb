"""
文档加载模块

支持 PDF、TXT、DOCX、MD、XLSX、PNG/JPG/WebP/BMP 多格式文档加载，
提供统一入口与动态分块策略选择。
"""
import os
from typing import Optional
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from docx import Document as DocxDocument
from utils.logger import logger
from utils.exceptions import DocumentLoadError
from core.strategy.chunk_strategy import get_document_splitter
from core.image_processor import describe_image
from core.multimodal import IMAGE_EXTENSIONS

def load_pdf(file_path: str) -> str:
    """加载 PDF 文件并提取纯文本内容"""
    try:
        loader = PyPDFLoader(file_path)
        pages = loader.load()
        text = "\n".join([page.page_content for page in pages])
        logger.info(f"成功加载PDF文件：{file_path}")
        return text
    except Exception as e:
        raise DocumentLoadError(f"PDF加载失败：{str(e)}")

def _load_text_with_fallback(file_path: str) -> str:
    """加载文本文件内容，utf-8 编码解析失败时自动回退 GBK

    Args:
        file_path: 文本文件路径

    Returns:
        文件文本内容
    """
    try:
        loader = TextLoader(file_path, encoding="utf-8")
        documents = loader.load()
    except UnicodeDecodeError:
        loader = TextLoader(file_path, encoding="gbk")
        documents = loader.load()
    return "\n".join([doc.page_content for doc in documents])

def load_txt(file_path: str) -> str:
    """加载 TXT 纯文本文件（utf-8 优先，失败回退 GBK）"""
    try:
        text = _load_text_with_fallback(file_path)
        logger.info(f"成功加载TXT文件：{file_path}")
        return text
    except Exception as e:
        raise DocumentLoadError(f"TXT加载失败：{str(e)}")

def load_docx(file_path: str) -> str:
    """加载 Word 文档并提取纯文本内容（含表格）"""
    try:
        doc = DocxDocument(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                full_text.append(row_text)
        text = "\n".join(full_text)
        logger.info(f"成功加载Word文档：{file_path}")
        return text
    except Exception as e:
        raise DocumentLoadError(f"Word加载失败：{str(e)}")

def load_md(file_path: str) -> str:
    """加载 Markdown 文档提取纯文本内容（utf-8 优先，失败回退 GBK）"""
    try:
        text = _load_text_with_fallback(file_path)
        logger.info(f"成功加载Markdown文件：{file_path}")
        return text
    except Exception as e:
        raise DocumentLoadError(f"Markdown加载失败：{str(e)}")

def load_xlsx(file_path: str) -> str:
    """加载 Excel 表格文档提取纯文本内容"""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        all_text = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            all_text.append(f"【工作表：{sheet_name}】")
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                row_text = row_text.strip()
                if row_text:
                    all_text.append(row_text)
        wb.close()
        text = "\n".join(all_text)
        logger.info(f"成功加载Excel文件：{file_path}")
        return text
    except Exception as e:
        raise DocumentLoadError(f"Excel加载失败：{str(e)}")


def load_image(file_path: str, config: Optional[dict] = None) -> str:
    """加载图片文件并调用视觉模型生成文字描述。"""
    try:
        text = describe_image(file_path, config=config)
        logger.info(f"成功加载图片文件：{file_path}")
        return text
    except Exception as e:
        raise DocumentLoadError(f"图片加载失败：{str(e)}")


def load_document(file_path: str, config: Optional[dict] = None) -> str:
    """通用文档加载入口，根据文件后缀自动识别格式并加载

    Args:
        file_path: 文档文件路径
        config: 全局配置字典，加载图片时透传给视觉模型识别，纯文本格式可省略

    Returns:
        提取的纯文本内容

    Raises:
        DocumentLoadError: 文件不存在或格式不支持时抛出
    """
    if not os.path.exists(file_path):
        raise DocumentLoadError(f"文件不存在：{file_path}")

    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return load_pdf(file_path)
    elif ext == ".txt":
        return load_txt(file_path)
    elif ext == ".docx":
        return load_docx(file_path)
    elif ext == ".md":
        return load_md(file_path)
    elif ext == ".xlsx":
        return load_xlsx(file_path)
    elif ext in IMAGE_EXTENSIONS:
        return load_image(file_path, config=config)
    else:
        raise DocumentLoadError(
            f"不支持的文件格式：{ext}，仅支持 .pdf/.txt/.docx/.md/.xlsx"
            " 以及 .png/.jpg/.jpeg/.webp/.bmp"
        )

def scan_directory_files(dir_path: str, support_exts: list[str]) -> list[str]:
    """遍历目录，筛选所有支持后缀的文档路径

    Args:
        dir_path: 目标目录路径
        support_exts: 支持的文件后缀列表，如 ['.pdf', '.txt']

    Returns:
        所有匹配文件的绝对路径列表
    """
    file_list = []
    for root, _, filenames in os.walk(dir_path):
        for filename in filenames:
            ext = os.path.splitext(filename)[1].lower()
            if ext in support_exts:
                full_path = os.path.join(root, filename)
                file_list.append(full_path)
    return file_list
